import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_table_of_contents(doc):
    """Inserts a dynamic table of contents that can be updated in Word."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(fldChar_end)

    doc.add_paragraph()  # Add space after TOC


def create_word_doc(content: str, filename: str) -> str:
    """Create a Word document from structured content with dynamic formatting and table of contents."""

    # Extract formatting
    font_match = re.search(r"\*\*Font:\*\*\s*([A-Za-z ]+),\s*(\d+)-point", content)
    font_name = font_match.group(1).strip() if font_match else "Times New Roman"
    font_size = int(font_match.group(2)) if font_match else 12

    spacing_match = re.search(r"\*\*Spacing:\*\*\s*(.+)", content)
    spacing_type = spacing_match.group(1).strip().lower() if spacing_match else "single"
    spacing_map = {"double": 2.0, "1.5": 1.5, "single": 1.0}
    line_spacing = spacing_map.get(spacing_type.lower(), 1.0)

    title_match = re.search(r"\*\*Assignment\s*\d+:\s*(.+?)\*\*", content)
    assignment_title = title_match.group(1).strip() if title_match else "Assignment"

    # Extract content body after "**Document:**" section
    document_start = re.search(r"\*\*Document:?\*\*[\s\S]+?\n", content)
    if not document_start:
        raise ValueError("Document content must follow '**Document:**' section with a blank line.")
    
    # Include the first line after **Document:** and the rest of the body
    body = content[document_start.end():].strip()

    # Initialize document
    doc = Document()

    # Set default font and size for the body text
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)

    # Title
    doc.add_heading(assignment_title, level=0)

    # Add TOC
    add_table_of_contents(doc)

    # Add the body content
    paragraphs = body.split('\n')

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Check if the paragraph is a heading (words wrapped in '**')
        if para.startswith("**") and para.endswith("**"):
            heading_text = para[2:-2]  # Remove the asterisks
            heading_paragraph = doc.add_heading(heading_text, level=1)
            # Set heading font size to 24pt
            for run in heading_paragraph.runs:
                run.font.size = Pt(24)
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = line_spacing
            # Handle bold text (e.g., **bold**)
            tokens = re.split(r"(\*\*[^\*]+\*\*)", para)
            for token in tokens:
                if token.startswith("**") and token.endswith("**"):
                    run = paragraph.add_run(token[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(token)

    # Save the document
    os.makedirs("generated_docs", exist_ok=True)
    filepath = os.path.join("generated_docs", filename)
    doc.save(filepath)
    return filepath
